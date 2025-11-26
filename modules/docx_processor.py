import os
import re
import shutil
from urllib.parse import urlparse, unquote

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT


def normalize_target(href, doc_saved_path=None, base_dir=None):
    href = href.strip()
    low = href.lower()
    p = urlparse(href)
    scheme = (p.scheme or '').lower()

    if scheme == 'mailto' or low.startswith('mailto:') or 'mailto:' in low:
        return 'email', href
    if 'skfgroup.sharepoint.com' in low:
        if 'document' in low:
            return 'document', href
        return 'internal', href
    if 'skf' in low:
        return 'internal', href
    if scheme in ('http', 'https', 'ftp') or href.startswith('//'):
        return 'external', href
    if scheme == 'file':
        path = unquote(p.path)
        if os.name == 'nt' and path.startswith('/') and len(path) > 2 and path[2] == ':':
            path = path.lstrip('/')
        norm = os.path.normpath(path)
        if base_dir:
            try:
                rel = os.path.relpath(norm, base_dir)
                return 'internal', rel.replace('\\', '/')
            except Exception:
                return 'internal', norm
        return 'internal', norm
    if os.name == 'nt' and len(href) > 2 and href[1] == ':' and href[2] in ('\\','/'):
        norm = os.path.normpath(href)
        if base_dir:
            try:
                rel = os.path.relpath(norm, base_dir)
                return 'internal', rel.replace('\\', '/')
            except Exception:
                return 'internal', norm
        return 'internal', norm
    if doc_saved_path:
        try:
            candidate = os.path.normpath(os.path.join(os.path.dirname(doc_saved_path), href))
            if base_dir:
                rel = os.path.relpath(candidate, base_dir)
                return 'internal', rel.replace('\\', '/')
            return 'internal', candidate
        except Exception:
            pass
    return 'unknown', href


def is_track_changes_on(doc):
    """Checks the settings.xml to see if Track Revisions is enabled."""
    try:
        settings = doc.settings.element
        track_revisions = settings.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}trackRevisions')
        return track_revisions is not None
    except Exception:
        return False

def get_links(doc, doc_path=None, base_dir=None):
    """Extracts hyperlinks and their types."""
    links = []
    rels = doc.part.rels
    
    # 1. Map Relationship IDs (rId) to actual URLs
    hyperlink_rels = {}
    for rel in rels.values():
        if rel.reltype == RT.HYPERLINK:
            hyperlink_rels[rel.rId] = rel.target_ref

    # 2. Scan document XML for hyperlink tags to get the clickable text
    for p in doc.paragraphs:
        for child in p._element:
            if child.tag.endswith('hyperlink'):
                rId = child.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                url = hyperlink_rels.get(rId)
                if url:
                    text = child.text if child.text else "[Image/Object]"
                    link_type, normalized = normalize_target(url, doc_saved_path=doc_path, base_dir=base_dir)
                    links.append({
                        "text": text,
                        "url": normalized,
                        "normalized": normalized,
                        "raw": url,
                        "type": link_type
                    })
    return links

def process_find_replace(file_path, find_text, replace_text=None, save_copy_path=None):
    """
    Scans for text. 
    Returns: match count, status, and list of text snippets (context).
    """
    # Use case-insensitive search
    if not find_text:
        return {'matches': 0, 'status': 'No find_text provided', 'snippets': []}

    doc = Document(file_path)
    matches_count = 0
    snippets = []
    found_urls = []
    found_texts = []
    found_urls = []
    found_texts = []
    found_urls = []
    found_texts = []

    # compile regex for case-insensitive literal match
    pattern = re.compile(re.escape(find_text), re.IGNORECASE)
    found_urls = []
    found_texts = []

    def process_element(paragraph_obj):
        nonlocal matches_count
        text = paragraph_obj.text or ''
        # count matches using regex
        found = pattern.findall(text)
        if not found:
            return
        matches_count += len(found)

        # Capture context snippet (truncate if too long)
        snippet = text.strip()
        if len(snippet) > 100:
            m = pattern.search(snippet)
            if m:
                idx = m.start()
            else:
                idx = snippet.lower().find(find_text.lower())
            start = max(0, idx - 40)
            end = min(len(snippet), idx + len(find_text) + 40)
            snippet = "..." + snippet[start:end] + "..."

        snippets.append(snippet)

        if replace_text is not None:
            # Replace all matches in a case-insensitive fashion
            # Note: this will not preserve original casing of matches
            new_text = pattern.sub(replace_text, text)
            paragraph_obj.text = new_text

    # Scan Paragraphs
    for p in doc.paragraphs:
        process_element(p)

    # Scan Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_element(p)

    copy_saved = None
    # If matches were found and a save path was provided, save a copy of the ORIGINAL file
    # even if we're only performing a find (no replacement). This ensures originals are kept.
    if matches_count > 0 and save_copy_path:
        try:
            os.makedirs(os.path.dirname(save_copy_path), exist_ok=True)
            shutil.copy2(file_path, save_copy_path)
            copy_saved = save_copy_path
        except Exception:
            # Don't fail the whole operation if copying fails; continue
            copy_saved = None

    if replace_text is not None and matches_count > 0:
        # After making a copy (if requested), save the modified document
        doc.save(file_path)
        status = 'Replaced & Saved'
    else:
        status = 'Found'

    result = {
        'matches': matches_count,
        'status': status,
        'snippets': snippets
    }
    if copy_saved:
        result['copy_path'] = copy_saved
    return result


def collect_links_for_files(file_paths, base_dir=None):
    """
    Returns a list of dicts: {'path': file_path, 'links': [...], 'error': Optional[str]}
    """
    results = []
    for path in file_paths:
        try:
            doc = Document(path)
            links = get_links(doc, doc_path=path, base_dir=base_dir or os.path.dirname(path))
            results.append({'path': path, 'links': links})
        except Exception as exc:
            results.append({'path': path, 'links': [], 'error': str(exc)})
    return results


def process_find_replace_bulk(file_paths, find_text, replace_text=None):
    """
    Runs find/replace across multiple files and aggregates results.
    """
    total_matches = 0
    per_file = []

    def _compute_save_path(path, base_dir, save_root):
        try:
            rel = os.path.relpath(path, base_dir)
        except Exception:
            rel = os.path.basename(path)
        return os.path.join(save_root, rel)

    # If caller provided a base_dir and requested copies, save under base_dir/bulk_found
    # The caller (app.bulk_find_replace) will pass base_dir when known.
    save_root = None
    # Caller may set attributes on function (we'll detect later in app call)

    for path in file_paths:
        try:
            # Determine save path per-file if save_root was provided by the caller via attribute
            save_copy_path = None
            # If caller attached a temporary attribute, use it (app will set save_root on this function)
            if hasattr(process_find_replace_bulk, '_save_root') and process_find_replace_bulk._save_root:
                save_copy_path = _compute_save_path(path, process_find_replace_bulk._base_dir, process_find_replace_bulk._save_root)

            result = process_find_replace(path, find_text, replace_text, save_copy_path)
            total_matches += result.get('matches', 0)
            per_file.append({'path': path, **result})
        except Exception as exc:
            per_file.append({'path': path, 'matches': 0, 'status': 'error', 'snippets': [], 'error': str(exc)})

    return {
        'total_matches': total_matches,
        'files': per_file,
        'mode': 'replace' if replace_text is not None else 'find'
    }


def process_links_find_replace(file_path, find_text, replace_text=None, target='both', save_copy_path=None):
    """Find/replace within hyperlinks of a document.

    - target: 'name' to target the link text, 'url' to target the hyperlink target, 'both' for both.
    - If replace_text is None, this will only search and report matches.
    Returns: dict with matches, status, snippets, and optional copy_path
    """
    if not find_text:
        return {'matches': 0, 'status': 'No find_text provided', 'snippets': []}

    doc = Document(file_path)
    matches_count = 0
    snippets = []
    found_urls = []
    found_texts = []

    pattern = re.compile(re.escape(find_text), re.IGNORECASE)

    def _process_hyperlink_element(child):
        nonlocal matches_count, snippets, found_urls, found_texts

        # Get relationship id and relationship object
        rId = child.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
        rel = None
        try:
            rel = doc.part.rels.get(rId) if rId else None
        except Exception:
            rel = None

        url = getattr(rel, 'target_ref', None) if rel is not None else None

        # Extract visible link text (concatenate all <w:t> text)
        parts = [t.text for t in child.iter() if t.tag.endswith('t') and t.text]
        link_text = ''.join(parts)

        # Search/replace link text
        if target in ('name', 'both') and link_text and pattern.search(link_text):
            found_count = len(pattern.findall(link_text))
            matches_count += found_count
            snippets.append(f"text: {link_text}")
            found_texts.append(link_text)
            if replace_text is not None:
                for t in child.iter():
                    if t.tag.endswith('t') and t.text:
                        t.text = pattern.sub(replace_text, t.text)

        # Search/replace URL
        if target in ('url', 'both') and url and pattern.search(url):
            found_count = len(pattern.findall(url))
            matches_count += found_count
            snippets.append(f"url: {url}")
            found_urls.append(url)
            if replace_text is not None and rel is not None:
                try:
                    # Replace the entire URL with the provided replacement for clarity (not just the matched substring)
                    new_target = replace_text
                    # Preserve the relationship type (external/internal) when re-writing the target
                    is_external = getattr(rel, 'is_external', True)
                    new_rId = doc.part.relate_to(new_target, RT.HYPERLINK, is_external=is_external)
                    child.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id', new_rId)
                    # Try to set the relationship's target if accessible
                    try:
                        new_rel = doc.part.rels.get(new_rId)
                        if new_rel is not None:
                            if hasattr(new_rel, 'target_ref'):
                                try:
                                    new_rel.target_ref = new_target
                                except Exception:
                                    pass
                            elif hasattr(new_rel, '_target'):
                                try:
                                    new_rel._target = new_target
                                except Exception:
                                    pass
                    except Exception:
                        pass
                    snippets.append(f"replaced-url: {rel.target_ref} -> {new_target} (rId={new_rId})")
                except Exception as exc:
                    snippets.append(f"replace-url-failed: {getattr(exc, 'args', exc)}")

    # Also handle field-based hyperlinks like: <w:instrText>HYPERLINK "https://..."</w:instrText>
    def _process_field_hyperlinks(paragraph_obj):
        nonlocal matches_count, snippets, found_urls, found_texts
        # iterate over instrText elements in the paragraph XML
        for instr in paragraph_obj._element.iter():
            if instr.tag.endswith('instrText') and instr.text and 'HYPERLINK' in instr.text:
                instr_text = instr.text
                # try to extract URL inside quotes or after HYPERLINK
                m = re.search(r'HYPERLINK\s+"([^"]+)"', instr_text)
                if not m:
                    m = re.search(r"HYPERLINK\s+'([^']+)'", instr_text)
                if not m:
                    m = re.search(r'HYPERLINK\s+([^\s]+)', instr_text)
                url = m.group(1) if m else None

                # visible/display text for the field is the paragraph text
                link_text = paragraph_obj.text or ''

                # Search/replace URL in field instruction
                if target in ('url', 'both') and url and pattern.search(url):
                    found_count = len(pattern.findall(url))
                    matches_count += found_count
                    snippets.append(f"field-url: {url}")
                    found_urls.append(url)
                    if replace_text is not None:
                        try:
                            # Replace the entire URL for field hyperlinks as well
                            new_target = replace_text
                            # replace within the instr text
                            new_instr = instr_text.replace(url, new_target)
                            instr.text = new_instr
                            snippets.append(f"replaced-field-url: {url} -> {new_target}")
                        except Exception as exc:
                            snippets.append(f"replace-field-url-failed: {getattr(exc, 'args', exc)}")

                # Search/replace visible text
                if target in ('name', 'both') and link_text and pattern.search(link_text):
                    found_count = len(pattern.findall(link_text))
                    matches_count += found_count
                    snippets.append(f"field-text: {link_text}")
                    found_texts.append(link_text)
                    if replace_text is not None:
                        # replace in runs to preserve run structure
                        for run in paragraph_obj.runs:
                            if run.text:
                                run.text = pattern.sub(replace_text, run.text)

    # Process paragraphs
    for p in doc.paragraphs:
        _process_field_hyperlinks(p)
        for child in p._element:
            if child.tag.endswith('hyperlink'):
                _process_hyperlink_element(child)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _process_field_hyperlinks(p)
                    for child in p._element:
                        if child.tag.endswith('hyperlink'):
                            _process_hyperlink_element(child)

    copy_saved = None
    if matches_count > 0 and save_copy_path:
        try:
            os.makedirs(os.path.dirname(save_copy_path), exist_ok=True)
            shutil.copy2(file_path, save_copy_path)
            copy_saved = save_copy_path
        except Exception:
            copy_saved = None

    if replace_text is not None and matches_count > 0:
        doc.save(file_path)
        status = 'Replaced & Saved'
    else:
        status = 'Found'

    result = {'matches': matches_count, 'status': status, 'snippets': snippets}
    if copy_saved:
        result['copy_path'] = copy_saved
    # include detected url/text lists for debugging/visibility
    if found_urls:
        result['found_urls'] = list(dict.fromkeys(found_urls))
    if found_texts:
        result['found_texts'] = list(dict.fromkeys(found_texts))
    result['did_replace'] = (replace_text is not None and matches_count > 0)
    return result



def process_links_find_replace_bulk(file_paths, find_text, replace_text=None, target='both'):
    """
    Bulk runner for link-targeted find/replace across files.
    Honors temporary attributes _save_root and _base_dir if set by caller.
    """
    total_matches = 0
    per_file = []

    def _compute_save_path(path, base_dir, save_root):
        try:
            rel = os.path.relpath(path, base_dir)
        except Exception:
            rel = os.path.basename(path)
        return os.path.join(save_root, rel)

    for path in file_paths:
        try:
            save_copy_path = None
            if hasattr(process_links_find_replace_bulk, '_save_root') and process_links_find_replace_bulk._save_root:
                save_copy_path = _compute_save_path(path, process_links_find_replace_bulk._base_dir, process_links_find_replace_bulk._save_root)

            # If the file itself lives under the save_root, skip it to avoid processing saved copies
            try:
                if hasattr(process_links_find_replace_bulk, '_save_root') and process_links_find_replace_bulk._save_root:
                    norm_save = os.path.normcase(os.path.abspath(process_links_find_replace_bulk._save_root))
                    norm_path = os.path.normcase(os.path.abspath(path))
                    if norm_path.startswith(norm_save + os.sep) or norm_path == norm_save:
                        # Skip files that are inside the save_root (these are previously saved copies)
                        continue
            except Exception:
                pass

            # First run a detection (find-only) without copying to see if there are matches.
            detect = process_links_find_replace(path, find_text, replace_text=None, target=target, save_copy_path=None)
            if detect.get('matches', 0) > 0 and replace_text is not None:
                # There are matches and we need to replace — perform actual replace and save copy if requested
                result = process_links_find_replace(path, find_text, replace_text, target=target, save_copy_path=save_copy_path)
            else:
                # No matches or this was a find-only operation — return detection result
                result = detect

            total_matches += result.get('matches', 0)
            per_file.append({'path': path, **result})
        except Exception as exc:
            per_file.append({'path': path, 'matches': 0, 'status': 'error', 'snippets': [], 'error': str(exc)})

    return {
        'total_matches': total_matches,
        'files': per_file,
        'mode': 'replace' if replace_text is not None else 'find',
        'target': target
    }
